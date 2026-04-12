"""
sirius_arquivos.py — Leitor universal de arquivos

Suporta: PDF, DOCX, XLSX, CSV, JSON, XML, imagens (OCR),
         áudio (metadados), código, texto simples e mais.

Uso:
    from sirius_arquivos import SiriusArquivos
    arq = SiriusArquivos()
    resultado = arq.ler("/caminho/para/arquivo.pdf")
    print(resultado.texto)
    print(resultado.resumo)
"""

import os
import sys
import re
import json

diretorio_src = os.path.dirname(os.path.abspath(__file__))
if diretorio_src not in sys.path:
    sys.path.insert(0, diretorio_src)


# ---------------------------------------------------------------------------
# Resultado padronizado
# ---------------------------------------------------------------------------

class ResultadoArquivo:
    def __init__(self, caminho: str):
        self.caminho   = caminho
        self.nome      = os.path.basename(caminho)
        self.extensao  = os.path.splitext(caminho)[1].lower()
        self.texto     = ""        # conteúdo extraído
        self.resumo    = ""        # resumo gerado
        self.metadados = {}        # info extra (páginas, autor, etc)
        self.erro      = None      # erro se falhou
        self.sucesso   = False

    def __str__(self):
        if self.erro:
            return f"[ERRO ao ler {self.nome}]: {self.erro}"
        preview = self.texto[:200] + "..." if len(self.texto) > 200 else self.texto
        return f"[{self.nome}]\n{preview}"


# ---------------------------------------------------------------------------
# Leitores por tipo
# ---------------------------------------------------------------------------

def _ler_txt(caminho: str) -> str:
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            with open(caminho, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


def _ler_pdf(caminho: str) -> tuple[str, dict]:
    texto    = ""
    metadados = {}
    try:
        import pdfplumber
        with pdfplumber.open(caminho) as pdf:
            metadados["paginas"] = len(pdf.pages)
            for pag in pdf.pages:
                t = pag.extract_text()
                if t:
                    texto += t + "\n"
        return texto, metadados
    except ImportError:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(caminho)
        metadados["paginas"] = len(reader.pages)
        if reader.metadata:
            metadados["autor"] = reader.metadata.get("/Author", "")
            metadados["titulo"] = reader.metadata.get("/Title", "")
        for pag in reader.pages:
            texto += pag.extract_text() or ""
        return texto, metadados
    except ImportError:
        pass

    return "", {"erro": "pip install pdfplumber ou pypdf"}


def _ler_docx(caminho: str) -> tuple[str, dict]:
    try:
        import docx
        doc   = docx.Document(caminho)
        texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        meta  = {
            "paragrafos": len(doc.paragraphs),
            "secoes":     len(doc.sections),
        }
        # Tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                texto += "\n" + " | ".join(c.text for c in linha.cells)
        return texto, meta
    except ImportError:
        return "", {"erro": "pip install python-docx"}


def _ler_xlsx(caminho: str) -> tuple[str, dict]:
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(caminho, read_only=True, data_only=True)
        texto = ""
        meta  = {"abas": wb.sheetnames}
        for nome_aba in wb.sheetnames:
            ws     = wb[nome_aba]
            texto += f"\n[Aba: {nome_aba}]\n"
            for linha in ws.iter_rows(values_only=True):
                texto += " | ".join(str(c) for c in linha if c is not None) + "\n"
        return texto, meta
    except ImportError:
        pass

    try:
        import pandas as pd
        dfs   = pd.read_excel(caminho, sheet_name=None)
        texto = ""
        for nome, df in dfs.items():
            texto += f"\n[Aba: {nome}]\n{df.to_string()}\n"
        return texto, {"abas": list(dfs.keys())}
    except ImportError:
        return "", {"erro": "pip install openpyxl pandas"}


def _ler_csv(caminho: str) -> tuple[str, dict]:
    try:
        import pandas as pd
        df    = pd.read_csv(caminho, encoding="utf-8", on_bad_lines="skip")
        meta  = {"linhas": len(df), "colunas": list(df.columns)}
        return df.to_string(index=False), meta
    except ImportError:
        # Fallback sem pandas
        import csv
        linhas = []
        with open(caminho, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                linhas.append(" | ".join(row))
        return "\n".join(linhas), {"linhas": len(linhas)}


def _ler_json(caminho: str) -> tuple[str, dict]:
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            dados = json.load(f)
        texto = json.dumps(dados, ensure_ascii=False, indent=2)
        meta  = {"tipo": type(dados).__name__}
        if isinstance(dados, list):
            meta["itens"] = len(dados)
        elif isinstance(dados, dict):
            meta["chaves"] = list(dados.keys())[:10]
        return texto, meta
    except Exception as e:
        return "", {"erro": str(e)}


def _ler_xml(caminho: str) -> tuple[str, dict]:
    try:
        import xml.etree.ElementTree as ET
        tree  = ET.parse(caminho)
        root  = tree.getroot()
        texto = ET.tostring(root, encoding="unicode")
        # Remove tags para obter só o texto
        texto_limpo = re.sub(r"<[^>]+>", " ", texto)
        texto_limpo = re.sub(r"\s+", " ", texto_limpo).strip()
        return texto_limpo, {"raiz": root.tag}
    except Exception as e:
        return "", {"erro": str(e)}


def _ler_imagem(caminho: str) -> tuple[str, dict]:
    """Tenta OCR na imagem com pytesseract ou easyocr."""
    meta  = {}
    texto = ""

    # Obtém dimensões
    try:
        from PIL import Image
        img  = Image.open(caminho)
        meta = {"largura": img.width, "altura": img.height, "modo": img.mode}
    except ImportError:
        pass

    # OCR com pytesseract (mais leve)
    try:
        import pytesseract
        from PIL import Image
        img   = Image.open(caminho)
        texto = pytesseract.image_to_string(img, lang="por+eng")
        meta["ocr"] = "pytesseract"
        return texto.strip(), meta
    except ImportError:
        pass
    except Exception as e:
        meta["erro_ocr"] = str(e)

    # OCR com easyocr (mais pesado mas mais preciso)
    try:
        import easyocr
        reader = easyocr.Reader(["pt", "en"], gpu=False)
        result = reader.readtext(caminho, detail=0)
        texto  = " ".join(result)
        meta["ocr"] = "easyocr"
        return texto, meta
    except ImportError:
        meta["aviso"] = "pip install pytesseract ou easyocr para OCR"

    return texto, meta


def _ler_audio(caminho: str) -> tuple[str, dict]:
    """Extrai metadados de arquivos de áudio (sem transcrição)."""
    meta = {}
    try:
        import mutagen
        audio = mutagen.File(caminho)
        if audio:
            meta = {
                "duracao_s": getattr(audio.info, "length", 0),
                "bitrate":   getattr(audio.info, "bitrate", 0),
                "tags":      {k: str(v) for k, v in (audio.tags or {}).items()},
            }
            texto = f"Arquivo de áudio: {os.path.basename(caminho)}\n"
            texto += f"Duração: {meta.get('duracao_s', 0):.1f}s\n"
            for k, v in meta.get("tags", {}).items():
                texto += f"{k}: {v}\n"
            return texto, meta
    except ImportError:
        meta["aviso"] = "pip install mutagen para metadados de áudio"
    return f"Arquivo de áudio: {os.path.basename(caminho)}", meta


def _ler_codigo(caminho: str) -> tuple[str, dict]:
    """Lê qualquer arquivo de código com análise básica."""
    texto = _ler_txt(caminho)
    linhas = texto.split("\n")
    meta = {
        "linhas":    len(linhas),
        "nao_vazias": sum(1 for l in linhas if l.strip()),
        "comentarios": sum(1 for l in linhas if l.strip().startswith(("#", "//", "/*", "*", "--"))),
    }
    # Detecta funções/classes (Python, JS, etc)
    funcoes = re.findall(r"(?:def|function|class|func)\s+(\w+)", texto)
    if funcoes:
        meta["funcoes_classes"] = funcoes[:20]
    return texto, meta


# ---------------------------------------------------------------------------
# Mapeamento de extensões
# ---------------------------------------------------------------------------

EXTENSOES_SUPORTADAS = {
    # Documentos
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".odt":  "docx",
    # Planilhas
    ".xlsx": "xlsx",
    ".xls":  "xlsx",
    ".ods":  "xlsx",
    ".csv":  "csv",
    ".tsv":  "csv",
    # Dados estruturados
    ".json": "json",
    ".xml":  "xml",
    ".yaml": "txt",
    ".yml":  "txt",
    ".toml": "txt",
    # Texto simples
    ".txt":  "txt",
    ".md":   "txt",
    ".rst":  "txt",
    ".log":  "txt",
    ".ini":  "txt",
    ".cfg":  "txt",
    ".env":  "txt",
    # Código
    ".py":   "codigo",
    ".js":   "codigo",
    ".ts":   "codigo",
    ".html": "codigo",
    ".css":  "codigo",
    ".java": "codigo",
    ".cpp":  "codigo",
    ".c":    "codigo",
    ".cs":   "codigo",
    ".go":   "codigo",
    ".rs":   "codigo",
    ".sql":  "codigo",
    ".sh":   "codigo",
    ".bat":  "codigo",
    # Imagens
    ".png":  "imagem",
    ".jpg":  "imagem",
    ".jpeg": "imagem",
    ".bmp":  "imagem",
    ".tiff": "imagem",
    ".webp": "imagem",
    # Áudio
    ".mp3":  "audio",
    ".wav":  "audio",
    ".flac": "audio",
    ".ogg":  "audio",
    ".m4a":  "audio",
}


# ---------------------------------------------------------------------------
# Classe principal
# ---------------------------------------------------------------------------

class SiriusArquivos:
    """
    Lê qualquer tipo de arquivo e retorna um ResultadoArquivo
    com o texto extraído, resumo e metadados.
    """

    def ler(self, caminho: str) -> ResultadoArquivo:
        """Lê o arquivo e retorna ResultadoArquivo."""
        resultado = ResultadoArquivo(caminho)

        if not os.path.exists(caminho):
            resultado.erro = f"Arquivo não encontrado: {caminho}"
            return resultado

        ext  = resultado.extensao
        tipo = EXTENSOES_SUPORTADAS.get(ext)

        if tipo is None:
            # Tenta como texto puro (último recurso)
            texto = _ler_txt(caminho)
            if texto:
                tipo = "txt"
            else:
                resultado.erro = f"Tipo de arquivo não suportado: {ext}"
                return resultado

        try:
            if tipo == "txt":
                resultado.texto     = _ler_txt(caminho)
                resultado.metadados = {"linhas": resultado.texto.count("\n")}

            elif tipo == "pdf":
                resultado.texto, resultado.metadados = _ler_pdf(caminho)

            elif tipo == "docx":
                resultado.texto, resultado.metadados = _ler_docx(caminho)

            elif tipo == "xlsx":
                resultado.texto, resultado.metadados = _ler_xlsx(caminho)

            elif tipo == "csv":
                resultado.texto, resultado.metadados = _ler_csv(caminho)

            elif tipo == "json":
                resultado.texto, resultado.metadados = _ler_json(caminho)

            elif tipo == "xml":
                resultado.texto, resultado.metadados = _ler_xml(caminho)

            elif tipo == "imagem":
                resultado.texto, resultado.metadados = _ler_imagem(caminho)

            elif tipo == "audio":
                resultado.texto, resultado.metadados = _ler_audio(caminho)

            elif tipo == "codigo":
                resultado.texto, resultado.metadados = _ler_codigo(caminho)

            resultado.sucesso = True
            resultado.resumo  = self._gerar_resumo(resultado)

        except Exception as e:
            resultado.erro = str(e)

        return resultado

    def _gerar_resumo(self, resultado: ResultadoArquivo) -> str:
        """Gera um resumo do arquivo para o Sirius apresentar."""
        texto = resultado.texto
        if not texto:
            return f"Arquivo {resultado.nome} lido, mas sem conteúdo extraível."

        # Resumo básico: primeiras 3 frases
        frases = re.split(r"[.!?]\s+", texto[:1000])
        frases = [f.strip() for f in frases if len(f.strip()) > 20][:3]
        preview = ". ".join(frases) + "." if frases else texto[:200]

        info = f"Arquivo '{resultado.nome}'"
        if "paginas" in resultado.metadados:
            info += f" ({resultado.metadados['paginas']} páginas)"
        elif "linhas" in resultado.metadados:
            info += f" ({resultado.metadados['linhas']} linhas)"
        elif "linhas" in resultado.metadados:
            info += f" ({resultado.metadados['linhas']} linhas de dados)"

        return f"{info}:\n{preview}"

    def ler_e_salvar_no_banco(self, caminho: str, memoria) -> ResultadoArquivo:
        """
        Lê o arquivo E salva o conteúdo no banco de treino do Sirius.
        Permite que o Sirius aprenda com o arquivo lido.
        """
        resultado = self.ler(caminho)
        if resultado.sucesso and resultado.texto:
            tema = os.path.splitext(resultado.nome)[0]
            memoria.salvar_estudo_autonomo(
                tema=tema,
                conteudo=resultado.texto[:3000],
                tags="arquivo_lido"
            )
            print(f"\033[92m[ARQUIVOS]: '{resultado.nome}' salvo no banco de treino.\033[0m")
        return resultado

    def listar_suportados(self) -> list[str]:
        """Retorna lista de extensões suportadas."""
        return sorted(EXTENSOES_SUPORTADAS.keys())

    def detectar_arquivo_no_comando(self, texto: str) -> str | None:
        """
        Tenta encontrar um caminho de arquivo mencionado no texto.
        Ex: 'lê o arquivo C:/Users/carlos/doc.pdf' → 'C:/Users/carlos/doc.pdf'
        """
        # Padrões de caminho Windows e Unix
        padroes = [
            r'[A-Za-z]:\\[^\s"\']+',                           # Windows: C:\...
            r'/(?:home|mnt|tmp|var|usr|opt)/[^\s"\']+',        # Unix absoluto
            r'"([^"]+\.[a-zA-Z]{2,5})"',                       # entre aspas duplas
            r"'([^']+\.[a-zA-Z]{2,5})'",                       # entre aspas simples
        ]
        for padrao in padroes:
            m = re.search(padrao, texto)
            if m:
                caminho = m.group(1) if m.lastindex else m.group()
                if os.path.exists(caminho):
                    return caminho

        # Tenta extensões conhecidas sem caminho completo
        ext_regex = "|".join(re.escape(e) for e in EXTENSOES_SUPORTADAS)
        m = re.search(r'(\S+(?:' + ext_regex + r'))', texto, re.IGNORECASE)
        if m:
            nome = m.group(1)
            # Procura nas pastas comuns
            pastas = [
                os.path.expanduser("~"),
                os.path.join(os.path.expanduser("~"), "Documents"),
                os.path.join(os.path.expanduser("~"), "Desktop"),
                os.path.join(os.path.expanduser("~"), "Downloads"),
            ]
            for pasta in pastas:
                caminho = os.path.join(pasta, nome)
                if os.path.exists(caminho):
                    return caminho

        return None


# ---------------------------------------------------------------------------
# Standalone
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Uso: python sirius_arquivos.py <caminho_do_arquivo>")
        sys.exit(1)

    arq = SiriusArquivos()
    res = arq.ler(sys.argv[1])
    print(f"\nSucesso: {res.sucesso}")
    print(f"Resumo: {res.resumo}")
    print(f"Metadados: {res.metadados}")
    if res.erro:
        print(f"Erro: {res.erro}")